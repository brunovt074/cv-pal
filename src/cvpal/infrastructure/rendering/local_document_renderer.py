"""Local DocumentRenderPort implementation - no agent, no API key needed.

`python-docx` (already a dependency) converts markdown straight to
`.docx`; `.pdf` goes through a temporary `.docx` plus `soffice --headless
--convert-to pdf` (LibreOffice is already installed on this machine - the
project previously assumed Anthropic Agent Skills were required for
document generation, which was never actually true here; see AGENTS.md
Decision Log).

Hyperlink preservation: `python-docx` writes proper `<w:hyperlink>`
elements with external relationships and Word-style blue/underline
formatting, so the .docx opens with clickable links in Word/Google
Docs/LibreOffice. LibreOffice's headless `--convert-to pdf` filter, on
the other hand, drops those external link relationships and emits a
PDF with the link text but no /Link annotations and no /URI actions -
verified on the system install (LibreOffice 24.2.7.2). To get
clickable links in the PDF, after soffice writes the PDF we read the
docx's relationships back out, locate each link's text in the PDF
with pdfplumber, and add /Link + /URI annotations with pypdf. Same
markdown in, clickable .docx AND .pdf out, no agent.
"""

from __future__ import annotations

import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

import pdfplumber
import pypdf
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pypdf.annotations import Link
from pypdf.generic import ArrayObject, NameObject, RectangleObject

from cvpal.domain.errors import DocumentRenderError
from cvpal.domain.generation.models import DocumentFormat

_SOFFICE_TIMEOUT_SECONDS = 60
_DEFAULT_FONT = "Arial"
_HYPERLINK_COLOR = "0563C1"
_STYLED_TEXT_STYLES = ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet")
_TRAILING_URL_PUNCTUATION = ".,;:"

_INLINE_PATTERN = re.compile(
    r"\*{1,2}(?P<bold>[^*\n]+?)\*{1,2}"
    r"|\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^\s)]+)\)"
    r"|(?P<bare_url>https?://[^\s)]+)"
)


def _apply_default_font(document: Document) -> None:
    for style_name in _STYLED_TEXT_STYLES:
        style = document.styles[style_name]
        style.font.name = _DEFAULT_FONT
        fonts = style.element.get_or_add_rPr().find(qn("w:rFonts"))
        for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if fonts.get(qn(theme_attr)) is not None:
                del fonts.attrib[qn(theme_attr)]
        fonts.set(qn("w:eastAsia"), _DEFAULT_FONT)
        fonts.set(qn("w:cs"), _DEFAULT_FONT)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), _DEFAULT_FONT)
    fonts.set(qn("w:hAnsi"), _DEFAULT_FONT)
    run_properties.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), _HYPERLINK_COLOR)
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def _add_markdown_paragraph(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    pos = 0
    for match in _INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])

        if match.group("bold") is not None:
            bold_run = paragraph.add_run(match.group("bold"))
            bold_run.bold = True
            pos = match.end()
        elif match.group("link_url") is not None:
            _add_hyperlink(paragraph, match.group("link_url"), match.group("link_text"))
            pos = match.end()
        else:
            url = match.group("bare_url")
            while url and url[-1] in _TRAILING_URL_PUNCTUATION:
                url = url[:-1]
            _add_hyperlink(paragraph, url, url)
            pos = match.start() + len(url)
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(markdown: str) -> Document:
    """Covers the markdown shape the tailoring/cv_pal prompts actually
    produce: headings, bold, bullets, plain paragraphs, and links (markdown
    `[text](url)` syntax or bare URLs). Not a general markdown renderer -
    an agent that wants something this doesn't cover can ask for the
    markdown format instead.
    """
    document = Document()
    _apply_default_font(document)
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            content = line.lstrip()[2:].strip()
            _add_markdown_paragraph(document, content, style="List Bullet")
        else:
            _add_markdown_paragraph(document, line)
    return document


def extract_docx_hyperlinks(docx_path: Path) -> list[tuple[str, str]]:
    """Read the hyperlink relationships + text out of a .docx file as
    (display_text, url) pairs, in document order. Used by the PDF
    pipeline to re-inject the clickable annotations LibreOffice's
    headless PDF export drops.
    """
    with zipfile.ZipFile(docx_path) as z:
        rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")
        document_xml = z.read("word/document.xml").decode("utf-8")

    url_by_id: dict[str, str] = {}
    for rel_id, target in re.findall(
        r'Id="(rId\d+)"[^>]*Target="([^"]+)"[^>]*TargetMode="External"', rels_xml
    ):
        url_by_id[rel_id] = target

    pairs: list[tuple[str, str]] = []
    for rel_id, inner in re.findall(
        r'<w:hyperlink[^>]*r:id="(rId\d+)"[^>]*>(.*?)</w:hyperlink>',
        document_xml,
        re.DOTALL,
    ):
        url = url_by_id.get(rel_id)
        if url is None:
            continue
        text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", inner))
        if text:
            pairs.append((text, url))
    return pairs


def _annotate_pdf_with_links(pdf_path: Path, link_data: list[tuple[str, str]]) -> None:
    """In-place: open the PDF, find each link text's bounding box via
    pdfplumber, and add a /Link annotation with a /URI action so the
    PDF actually becomes clickable (LibreOffice's headless export
    drops these on the way in — and even when it does preserve some
    on a given run, the result is inconsistent, so we strip any
    /Link annotations we find on each page and re-add our own).
    """
    if not link_data:
        return

    reader = pypdf.PdfReader(str(pdf_path))
    writer = pypdf.PdfWriter()
    writer.append_pages_from_reader(reader)

    for page in writer.pages:
        if "/Annots" not in page:
            continue
        page[NameObject("/Annots")] = ArrayObject(
            annot
            for annot in page["/Annots"]
            if annot.get_object().get("/Subtype") != "/Link"
        )

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_index, page in enumerate(pdf.pages):
            page_height = page.height
            for text, url in link_data:
                matches = page.search(text)
                if not matches:
                    continue
                for rect_match in matches:
                    x0, top, x1, bottom = (
                        rect_match["x0"],
                        rect_match["top"],
                        rect_match["x1"],
                        rect_match["bottom"],
                    )
                    y0 = page_height - bottom
                    y1 = page_height - top
                    writer.add_annotation(
                        page_number=page_index,
                        annotation=Link(
                            rect=RectangleObject([x0, y0, x1, y1]),
                            url=url,
                        ),
                    )

    with pdf_path.open("wb") as fh:
        writer.write(fh)


class LocalDocumentRenderer:
    def render(self, content: str, *, document_format: DocumentFormat, output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if document_format == DocumentFormat.MARKDOWN:
            output_path.write_text(content)
            return output_path
        if document_format == DocumentFormat.DOCX:
            markdown_to_docx(content).save(output_path)
            return output_path
        if document_format == DocumentFormat.PDF:
            return self._render_pdf(content, output_path)
        raise DocumentRenderError(f"Unsupported document format: {document_format}")

    def _render_pdf(self, content: str, output_path: Path) -> Path:
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            docx_path = tmp_path / "source.docx"
            document = markdown_to_docx(content)
            document.save(docx_path)
            link_data = extract_docx_hyperlinks(docx_path)

            profile_dir = tmp_path / "soffice-profile"
            try:
                result = subprocess.run(
                    [
                        "soffice",
                        "--headless",
                        f"-env:UserInstallation=file://{profile_dir}",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(tmp_path),
                        str(docx_path),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=_SOFFICE_TIMEOUT_SECONDS,
                )
            except subprocess.TimeoutExpired as exc:
                raise DocumentRenderError(
                    f"soffice timed out after {_SOFFICE_TIMEOUT_SECONDS}s converting to pdf"
                ) from exc
            except FileNotFoundError as exc:
                raise DocumentRenderError("soffice binary not found - is LibreOffice installed?") from exc

            if result.returncode != 0:
                raise DocumentRenderError(f"soffice failed converting to pdf: {result.stderr[:2000]}")

            produced = tmp_path / "source.pdf"
            if not produced.exists():
                raise DocumentRenderError("soffice did not produce a PDF output")

            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_bytes(produced.read_bytes())
            _annotate_pdf_with_links(output_path, link_data)
            return output_path
