from __future__ import annotations

from pathlib import Path

import docx


def _extract_hyperlinks(document: docx.Document) -> list[str]:
    """Plain-text extraction loses hyperlink targets (e.g. a "LinkedIn" link
    renders as the word "LinkedIn", not its URL) - relationship targets are
    separate from the paragraph text, so we pull them separately.
    """
    urls: list[str] = []
    for rel in document.part.rels.values():
        if "hyperlink" in rel.reltype and rel.target_ref not in urls:
            urls.append(rel.target_ref)
    return urls


def extract_docx_text(path: Path) -> str:
    document = docx.Document(str(path))
    parts: list[str] = []

    for para in document.paragraphs:
        parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    urls = _extract_hyperlinks(document)
    if urls:
        parts.append("")
        parts.append("Links:")
        parts.extend(urls)

    return "\n".join(parts)
