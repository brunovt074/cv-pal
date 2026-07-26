import shutil
from pathlib import Path

import pypdf
import pytest
from docx import Document
from docx.oxml.ns import qn

from cvpal.domain.errors import DocumentRenderError
from cvpal.domain.generation.models import DocumentFormat
from cvpal.infrastructure.rendering.local_document_renderer import (
    LocalDocumentRenderer,
    extract_docx_hyperlinks,
    markdown_to_docx,
)


def _hyperlinks(paragraph):
    links = []
    for hyperlink in paragraph._p.findall(qn("w:hyperlink")):
        r_id = hyperlink.get(qn("r:id"))
        url = paragraph.part.rels[r_id].target_ref
        text = "".join(node.text or "" for node in hyperlink.iter(qn("w:t")))
        links.append((text, url))
    return links

_SAMPLE_MARKDOWN = """\
# Bruno Vargas Tettamanti

## Summary

Backend developer with **Java** and Spring Boot experience.

## Experience

### Acme — Developer

- Built REST APIs
- Led a **migration** project
"""

_SOFFICE_AVAILABLE = shutil.which("soffice") is not None


def test_markdown_to_docx_maps_headings():
    document = markdown_to_docx(_SAMPLE_MARKDOWN)
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert "Bruno Vargas Tettamanti" in headings
    assert "Summary" in headings
    assert "Experience" in headings
    assert "Acme — Developer" in headings


def test_markdown_to_docx_keeps_bullet_content():
    document = markdown_to_docx(_SAMPLE_MARKDOWN)
    bullets = [p.text for p in document.paragraphs if p.style.name == "List Bullet"]
    assert "Built REST APIs" in bullets
    assert any("migration" in b for b in bullets)


def test_markdown_to_docx_marks_bold_runs():
    document = markdown_to_docx("Backend developer with **Java** experience.")
    paragraph = document.paragraphs[0]
    bold_runs = [r for r in paragraph.runs if r.bold]
    assert any(r.text == "Java" for r in bold_runs)


def test_markdown_to_docx_treats_single_asterisk_as_bold():
    document = markdown_to_docx("Dates *(09/2023 – 02/2024, Remote)* matter.")
    paragraph = document.paragraphs[0]
    bold_runs = [r for r in paragraph.runs if r.bold]
    assert any("(09/2023 – 02/2024, Remote)" in r.text for r in bold_runs)


def test_markdown_to_docx_does_not_match_bold_across_asterisk_pairs():
    document = markdown_to_docx("*alpha* and *beta* are separate.")
    paragraph = document.paragraphs[0]
    bold_runs = [r for r in paragraph.runs if r.bold]
    bold_texts = [r.text for r in bold_runs]
    assert "alpha" in bold_texts
    assert "beta" in bold_texts


def test_markdown_to_docx_uses_arial_as_default_font():
    document = markdown_to_docx(_SAMPLE_MARKDOWN)
    assert document.styles["Normal"].font.name == "Arial"
    assert document.styles["Heading 1"].font.name == "Arial"
    assert document.styles["Heading 2"].font.name == "Arial"
    assert document.styles["List Bullet"].font.name == "Arial"


def test_markdown_to_docx_renders_markdown_links_as_real_hyperlinks():
    document = markdown_to_docx(
        "Reach me on [LinkedIn](https://www.linkedin.com/in/bruno-vargas-tettamanti-dev)."
    )
    paragraph = document.paragraphs[0]
    links = _hyperlinks(paragraph)
    assert links == [("LinkedIn", "https://www.linkedin.com/in/bruno-vargas-tettamanti-dev")]
    assert "[LinkedIn]" not in paragraph.text
    assert paragraph.text.endswith(".")


def test_markdown_to_docx_renders_bare_urls_as_real_hyperlinks():
    document = markdown_to_docx("GitHub: https://github.com/brunovt074, always up to date.")
    paragraph = document.paragraphs[0]
    links = _hyperlinks(paragraph)
    assert links == [("https://github.com/brunovt074", "https://github.com/brunovt074")]
    assert paragraph.text.endswith("always up to date.")


def test_render_markdown_format_writes_content_as_is(tmp_path):
    renderer = LocalDocumentRenderer()
    output = tmp_path / "cv.md"
    result = renderer.render(_SAMPLE_MARKDOWN, document_format=DocumentFormat.MARKDOWN, output_path=output)
    assert result == output
    assert output.read_text() == _SAMPLE_MARKDOWN


def test_render_docx_format_produces_openable_document(tmp_path):
    renderer = LocalDocumentRenderer()
    output = tmp_path / "cv.docx"
    result = renderer.render(_SAMPLE_MARKDOWN, document_format=DocumentFormat.DOCX, output_path=output)
    assert result == output
    assert output.exists()
    reopened = Document(str(output))
    assert any("Bruno Vargas Tettamanti" in p.text for p in reopened.paragraphs)


@pytest.mark.skipif(not _SOFFICE_AVAILABLE, reason="soffice not installed")
def test_render_pdf_format_produces_a_real_pdf(tmp_path):
    renderer = LocalDocumentRenderer()
    output = tmp_path / "cv.pdf"
    result = renderer.render(_SAMPLE_MARKDOWN, document_format=DocumentFormat.PDF, output_path=output)
    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0
    assert output.read_bytes()[:4] == b"%PDF"


def test_extract_docx_hyperlinks_returns_text_url_pairs():
    markdown = (
        "Reach me on [LinkedIn](https://www.linkedin.com/in/bruno-vargas-tettamanti-dev) "
        "or check [GitHub](https://github.com/brunovt074)."
    )
    document = markdown_to_docx(markdown)
    docx_path = Path("/tmp/_extract_docx_hyperlinks.docx")
    document.save(docx_path)

    pairs = extract_docx_hyperlinks(docx_path)
    assert ("LinkedIn", "https://www.linkedin.com/in/bruno-vargas-tettamanti-dev") in pairs
    assert ("GitHub", "https://github.com/brunovt074") in pairs


def test_extract_docx_hyperlinks_ignores_plain_paragraphs():
    document = markdown_to_docx("Just a plain paragraph with no links.")
    docx_path = Path("/tmp/_extract_docx_hyperlinks_plain.docx")
    document.save(docx_path)

    assert extract_docx_hyperlinks(docx_path) == []


@pytest.mark.skipif(not _SOFFICE_AVAILABLE, reason="soffice not installed")
def test_render_pdf_preserves_hyperlinks_as_clickable_annotations(tmp_path):
    renderer = LocalDocumentRenderer()
    output = tmp_path / "cv.pdf"
    markdown = (
        "Contact: [LinkedIn](https://www.linkedin.com/in/bruno-vargas-tettamanti-dev) "
        "and [GitHub](https://github.com/brunovt074).\n"
    )
    renderer.render(markdown, document_format=DocumentFormat.PDF, output_path=output)

    reader = pypdf.PdfReader(str(output))
    page = reader.pages[0]
    link_uris = sorted(
        annot.get_object()["/A"]["/URI"]
        for annot in page["/Annots"]
        if annot.get_object().get("/Subtype") == "/Link"
        and annot.get_object().get("/A", {}).get("/S") == "/URI"
    )
    assert link_uris == [
        "https://github.com/brunovt074",
        "https://www.linkedin.com/in/bruno-vargas-tettamanti-dev",
    ]


def test_render_unsupported_format_raises(tmp_path):
    renderer = LocalDocumentRenderer()
    with pytest.raises(DocumentRenderError):
        renderer.render("x", document_format="not-a-format", output_path=tmp_path / "x")
